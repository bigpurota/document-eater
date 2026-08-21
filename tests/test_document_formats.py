from __future__ import annotations

import json

import pytest
from defusedxml.common import EntitiesForbidden
from docx import Document as WordDocument
from openpyxl import Workbook

from document_eater.audit import audit_corpus
from document_eater.index import index_artifacts, search
from document_eater.ingest import discover_documents, ingest_document


def _word(path) -> None:
    document = WordDocument()
    document.add_heading("Contract requirements", level=1)
    document.add_paragraph("The supplier must submit the signed acceptance certificate.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Deliverable"
    table.cell(0, 1).text = "Deadline"
    table.cell(1, 0).text = "Final report"
    table.cell(1, 1).text = "2026-09-01"
    document.save(path)


def _excel(path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Evidence"
    sheet.append(["Deliverable", "Status", "Check"])
    sheet.append(["Signed acceptance certificate", "Submitted", "=1+1"])
    second = workbook.create_sheet("Schedule")
    second.append(["Final report", "2026-09-01"])
    workbook.save(path)


def test_docx_preserves_headings_tables_and_locations(tmp_path):
    source = tmp_path / "requirements.docx"
    _word(source)

    destination = ingest_document(source, tmp_path / "artifacts")
    document = json.loads((destination / "document.json").read_text())
    manifest = json.loads((destination / "manifest.json").read_text())

    assert document["format"] == "docx"
    assert document["pages"][0]["kind"] == "document"
    assert document["pages"][0]["blocks"][0]["role"] == "heading"
    assert any(block["attrs"].get("table") == 1 for block in document["pages"][0]["blocks"])
    assert any(
        "paragraph" in block["attrs"]["location"] for block in document["pages"][0]["blocks"]
    )
    assert manifest["source_format"] == "docx"
    assert manifest["source_copied"] is False


def test_xlsx_preserves_sheets_cells_and_formulas_without_execution(tmp_path):
    source = tmp_path / "evidence.xlsx"
    _excel(source)

    destination = ingest_document(source, tmp_path / "artifacts")
    document = json.loads((destination / "document.json").read_text())
    manifest = json.loads((destination / "manifest.json").read_text())

    assert document["format"] == "xlsx"
    assert [page["label"] for page in document["pages"]] == ["Evidence", "Schedule"]
    evidence_blocks = document["pages"][0]["blocks"]
    formula_cells = [
        cell
        for block in evidence_blocks
        for cell in block["attrs"].get("cells", [])
        if cell["formula"]
    ]
    assert formula_cells[0]["coordinate"] == "C2"
    assert formula_cells[0]["formula"] == "=1+1"
    assert manifest["settings"]["formula_policy"] == "preserve_without_execution"

    index_artifacts(tmp_path / "artifacts", tmp_path / "index.sqlite3")
    hits = search(tmp_path / "index.sqlite3", "acceptance certificate submitted")
    assert hits[0].location_start == "sheet Evidence"
    assert hits[0].location_end.startswith("Evidence!")


def test_txt_csv_and_xml_join_the_same_index(tmp_path):
    source = tmp_path / "sources"
    source.mkdir()
    (source / "notes.TXT").write_bytes(
        "Требования\n\nИсполнитель должен представить отчет.".encode("cp1251")
    )
    (source / "status.csv").write_text("item;status\nотчет;представлен\n", encoding="utf-8")
    (source / "readme.md").write_text(
        "# Дополнительные условия\n\nПодрядчик обязан сохранить журнал.", encoding="utf-8"
    )
    (source / "contract.xml").write_text(
        "<contract><clause id='7'>Заказчик обязан утвердить результат.</clause></contract>",
        encoding="utf-8",
    )

    documents = discover_documents(source)
    assert [path.suffix.casefold() for path in documents] == [".xml", ".txt", ".md", ".csv"]
    artifacts = tmp_path / "artifacts"
    for document in documents:
        ingest_document(document, artifacts)
    stats = index_artifacts(artifacts, tmp_path / "index.sqlite3")

    assert stats["documents"] == 4
    assert search(tmp_path / "index.sqlite3", "представить отчет")
    assert search(tmp_path / "index.sqlite3", "заказчик утвердить")
    assert search(tmp_path / "index.sqlite3", "подрядчик журнал")


def test_mixed_word_excel_and_text_corpus_runs_end_to_end(tmp_path):
    source = tmp_path / "private-documents"
    source.mkdir()
    _word(source / "requirements.docx")
    _excel(source / "evidence.xlsx")
    (source / "notes.txt").write_text(
        "The project manager shall archive the final report.", encoding="utf-8"
    )

    run = tmp_path / "run"
    report = audit_corpus(source, run, retrieval_mode="lexical")
    manifest = json.loads((run / "run-manifest.json").read_text())

    assert len(report.items) == 2
    assert {item.requirement.filename for item in report.items} == {
        "notes.txt",
        "requirements.docx",
    }
    assert manifest["document_count"] == 3
    assert manifest["document_counts"]["docx"] == 1
    assert manifest["document_counts"]["xlsx"] == 1
    assert manifest["document_counts"]["txt"] == 1
    assert (run / "report.html").is_file()


def test_generated_workspaces_are_never_rediscovered_as_source_documents(tmp_path):
    source = tmp_path / "case"
    source.mkdir()
    original = source / "requirements.txt"
    original.write_text("The supplier must submit the final report.", encoding="utf-8")
    current_workspace = source / ".document-eater-workspace"
    current_workspace.mkdir()
    (current_workspace / ".document-eater-workspace").write_text("generated\n")
    (current_workspace / "requirements.csv").write_text("generated,result\n")
    legacy_workspace = source / "audit-run"
    legacy_workspace.mkdir()
    (legacy_workspace / "run-manifest.json").write_text("{}")
    (legacy_workspace / "requirements.csv").write_text("old,result\n")

    assert discover_documents(source) == [original.resolve()]


def test_unchanged_nested_workspace_reuses_the_completed_audit(tmp_path):
    source = tmp_path / "case"
    source.mkdir()
    (source / "requirements.txt").write_text(
        "The supplier must submit the final report.", encoding="utf-8"
    )
    workspace = source / ".document-eater-workspace"

    first = audit_corpus(source, workspace, retrieval_mode="lexical")
    first_manifest = (workspace / "run-manifest.json").read_text(encoding="utf-8")
    second = audit_corpus(source, workspace, retrieval_mode="lexical")

    assert first.reused is False
    assert second.reused is True
    assert (workspace / "run-manifest.json").read_text(encoding="utf-8") == first_manifest
    assert len(second.items) == 1

    (source / "requirements.txt").write_text(
        "The supplier must submit the final report.\nThe client shall approve it.",
        encoding="utf-8",
    )
    changed = audit_corpus(source, workspace, retrieval_mode="lexical")

    assert changed.reused is False
    assert len(changed.items) == 2
    assert (workspace / "run-manifest.json").read_text(encoding="utf-8") != first_manifest


def test_xml_external_entity_is_rejected(tmp_path):
    source = tmp_path / "unsafe.xml"
    source.write_text(
        '<!DOCTYPE x [<!ENTITY secret SYSTEM "file:///etc/passwd">]><x>&secret;</x>',
        encoding="utf-8",
    )

    with pytest.raises(EntitiesForbidden):
        ingest_document(source, tmp_path / "artifacts")
