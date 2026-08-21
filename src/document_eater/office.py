from __future__ import annotations

import re
from datetime import date, datetime, time
from pathlib import Path
from typing import Any

from docx import Document as OpenDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

from .artifacts import source_hash, write_document_artifacts
from .models import BBox, Block, Document, Page

_HEADING_STYLE = re.compile(r"^(?:heading|заголовок)\s*\d*", re.IGNORECASE)
_MAX_NONEMPTY_XLSX_CELLS = 250_000


def _outline_level(paragraph: Paragraph) -> int | None:
    properties = paragraph._p.pPr  # noqa: SLF001 - python-docx exposes no public outline API
    outline = properties.outlineLvl if properties is not None else None
    if outline is None or outline.val is None:
        return None
    return int(outline.val)


def _paragraph_font_size(paragraph: Paragraph) -> float | None:
    sizes = [run.font.size.pt for run in paragraph.runs if run.font.size]
    if not sizes and paragraph.style and paragraph.style.font.size:
        sizes = [paragraph.style.font.size.pt]
    return round(sum(sizes) / len(sizes), 2) if sizes else None


def _table_row_text(cells: list[str], headers: list[str]) -> str:
    parts = []
    for index, value in enumerate(cells, 1):
        if not value:
            continue
        header = headers[index - 1] if index <= len(headers) else ""
        parts.append(f"{header}: {value}" if header and header != value else value)
    return " | ".join(parts)


def ingest_docx(path: str | Path, output: str | Path) -> Path:
    source = Path(path).expanduser().resolve()
    if source.suffix.casefold() != ".docx" or not source.is_file():
        raise ValueError(f"Expected an existing DOCX file, got: {source}")

    sha256 = source_hash(source)
    doc_id = sha256[:16]
    opened = OpenDocument(source)
    blocks: list[Block] = []
    paragraph_number = 0
    table_number = 0

    for item in opened.iter_inner_content():
        if isinstance(item, Paragraph):
            text = "\n".join(line.strip() for line in item.text.splitlines() if line.strip())
            if not text:
                continue
            paragraph_number += 1
            style = item.style.name if item.style else ""
            heading = bool(_HEADING_STYLE.match(style)) or _outline_level(item) is not None
            order = len(blocks) + 1
            blocks.append(
                Block(
                    id=f"{doc_id}:p1:b{order}",
                    page=1,
                    order=order,
                    text=text,
                    bbox=None,
                    source="native",
                    font_size=_paragraph_font_size(item),
                    role="heading" if heading else "body",
                    attrs={
                        "location": f"paragraph {paragraph_number}",
                        "paragraph": paragraph_number,
                        "style": style,
                    },
                )
            )
            continue

        if isinstance(item, Table):
            table_number += 1
            rows = [[" ".join(cell.text.split()) for cell in row.cells] for row in item.rows]
            headers = rows[0] if rows else []
            for row_number, cells in enumerate(rows, 1):
                text = _table_row_text(cells, headers if row_number > 1 else [])
                if not text:
                    continue
                order = len(blocks) + 1
                blocks.append(
                    Block(
                        id=f"{doc_id}:p1:b{order}",
                        page=1,
                        order=order,
                        text=text,
                        bbox=None,
                        source="native",
                        attrs={
                            "location": f"table {table_number}, row {row_number}",
                            "table": table_number,
                            "row": row_number,
                            "cells": cells,
                        },
                    )
                )

    page = Page(
        number=1,
        width=0.0,
        height=0.0,
        source="native",
        native_char_count=sum(len("".join(block.text.split())) for block in blocks),
        blocks=blocks,
        kind="document",
        label="Word document",
    )
    document = Document(doc_id, source.name, sha256, [page], format="docx")
    return write_document_artifacts(
        document,
        source,
        output,
        settings={"parser": "python-docx", "pagination": "logical"},
    )


def _cell_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat(sep=" ")
    if isinstance(value, (date, time)):
        return value.isoformat()
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    return str(value)


def ingest_xlsx(path: str | Path, output: str | Path) -> Path:
    source = Path(path).expanduser().resolve()
    if source.suffix.casefold() != ".xlsx" or not source.is_file():
        raise ValueError(f"Expected an existing XLSX file, got: {source}")

    sha256 = source_hash(source)
    doc_id = sha256[:16]
    formulas = load_workbook(source, read_only=True, data_only=False, keep_links=False)
    values = load_workbook(source, read_only=True, data_only=True, keep_links=False)
    pages: list[Page] = []
    nonempty_cells = 0

    try:
        for sheet_number, (sheet, value_sheet) in enumerate(
            zip(formulas.worksheets, values.worksheets, strict=True), 1
        ):
            blocks = [
                Block(
                    id=f"{doc_id}:p{sheet_number}:b1",
                    page=sheet_number,
                    order=1,
                    text=f"Worksheet: {sheet.title}",
                    bbox=None,
                    source="native",
                    role="heading",
                    attrs={"location": f"sheet {sheet.title}", "sheet": sheet.title},
                )
            ]
            headers: dict[int, str] = {}
            for formula_row, value_row in zip(
                sheet.iter_rows(), value_sheet.iter_rows(), strict=True
            ):
                cell_records = []
                columns = []
                for formula_cell, value_cell in zip(formula_row, value_row, strict=True):
                    if formula_cell.value is None and value_cell.value is None:
                        continue
                    nonempty_cells += 1
                    if nonempty_cells > _MAX_NONEMPTY_XLSX_CELLS:
                        raise ValueError(
                            f"Workbook exceeds {_MAX_NONEMPTY_XLSX_CELLS} non-empty cells: {source}"
                        )
                    coordinate = formula_cell.coordinate
                    formula = str(formula_cell.value) if formula_cell.data_type == "f" else None
                    cached = _cell_value(value_cell.value) if formula else ""
                    value = _cell_value(formula_cell.value)
                    display = value
                    if formula and cached and cached != value:
                        display = f"{formula} => {cached}"
                    columns.append(formula_cell.column)
                    cell_records.append(
                        {
                            "column": formula_cell.column,
                            "coordinate": coordinate,
                            "value": value,
                            "display": display,
                            "formula": formula,
                            "cached_value": cached or None,
                            "number_format": formula_cell.number_format,
                        }
                    )
                if not cell_records:
                    continue
                is_header = (
                    not headers
                    and len(cell_records) >= 2
                    and all(not record["formula"] and record["value"] for record in cell_records)
                )
                if is_header:
                    headers = {
                        int(record["column"]): str(record["value"]) for record in cell_records
                    }
                rendered = []
                for record in cell_records:
                    label = headers.get(int(record["column"]))
                    prefix = str(record["coordinate"])
                    if label and not is_header:
                        prefix += f" ({label})"
                    rendered.append(f"{prefix}: {record['display']}")
                    record.pop("display")
                    record.pop("column")
                row_number = formula_row[0].row
                order = len(blocks) + 1
                blocks.append(
                    Block(
                        id=f"{doc_id}:p{sheet_number}:b{order}",
                        page=sheet_number,
                        order=order,
                        text=" | ".join(rendered),
                        bbox=BBox(
                            float(min(columns)),
                            float(row_number),
                            float(max(columns)),
                            float(row_number),
                        ),
                        source="native",
                        attrs={
                            "location": (
                                f"{sheet.title}!{get_column_letter(min(columns))}{row_number}:"
                                f"{get_column_letter(max(columns))}{row_number}"
                            ),
                            "sheet": sheet.title,
                            "row": row_number,
                            "cells": cell_records,
                        },
                    )
                )
            pages.append(
                Page(
                    number=sheet_number,
                    width=float(sheet.max_column or 0),
                    height=float(sheet.max_row or 0),
                    source="native",
                    native_char_count=sum(len("".join(block.text.split())) for block in blocks),
                    blocks=blocks,
                    kind="sheet",
                    label=sheet.title,
                    attrs={"visibility": sheet.sheet_state},
                )
            )
    finally:
        formulas.close()
        values.close()

    document = Document(doc_id, source.name, sha256, pages, format="xlsx")
    return write_document_artifacts(
        document,
        source,
        output,
        settings={
            "parser": "openpyxl",
            "formula_policy": "preserve_without_execution",
            "max_nonempty_cells": _MAX_NONEMPTY_XLSX_CELLS,
        },
    )
